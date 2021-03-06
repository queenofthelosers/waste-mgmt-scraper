import sys
import requests
from bs4 import BeautifulSoup
import docx

n = len(sys.argv)

if n==1 :
    print("Please enter the url")
elif n==2 :
    url = sys.argv[1]
    r = requests.get(url)
    htmlContent = r.content
    soup = BeautifulSoup(htmlContent, 'html.parser')
    #print(soup.prettify())
    module_title = soup.find('div',{'class':'video-link-wrapper current-video'}).find('span',{'class':'video-name'})
    myDoc = docx.Document("notes.docx")
    myDoc.add_heading(module_title, 2)
    for paragraph in soup.find_all('p',{'class' : "rc-plain-text-transcript body-1-text"}):
        para = str(paragraph.text)
        myDoc.add_paragraph(para)
    myDoc.save("notes.docx")
else :
    print("Follow format please")
